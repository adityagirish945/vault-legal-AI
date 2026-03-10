"""
Vault PropTech - Property Legal Assistant
User-facing chat interface powered by RAG + Gemini
"""

# SQLite3 Patch for Streamlit Cloud (Streamlit uses Debian which has old sqlite3)
try:
    __import__('pysqlite3')
    import sys
    sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')
except (ImportError, KeyError):
    pass

import os
import streamlit as st
from llm import ask, get_gemini_client
from query import query_kb, format_context_for_llm
from firebase_chat import save_chat, load_chats, load_chat, delete_chat, save_draft, get_draft
from vault_auth import check_auth, get_auth_config, get_login_url, clear_user_cookie
from drive_manager import upload_file, extract_text_with_gemini, export_to_google_doc, list_user_files
import uuid
import io
import re
import markdown as md_lib
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

# Page config
st.set_page_config(
    page_title="Vault PropTech Legal Assistant",
    layout="wide",
)

# ── Brand Assets ────────────────────────────────────────────────────────────
VAULT_LOGO_SVG = '''<svg width="84" height="30" viewBox="0 0 84 30" fill="none" xmlns="http://www.w3.org/2000/svg">
  <path d="M77.8514 11.5081H74.8991C74.7528 11.5081 74.6797 11.4425 74.6797 11.3114V10.1967C74.6797 10.0656 74.7528 10 74.8991 10H81.9008C82.6854 10 83.0777 10.3115 83.0777 10.9344V11.2786C83.0777 11.3332 83.0511 11.3879 82.9979 11.4425C82.958 11.4862 82.8915 11.5081 82.7984 11.5081H79.8262V18.9338C79.8262 19.054 79.7597 19.1141 79.6267 19.1141H78.0907C77.9312 19.1141 77.8514 19.054 77.8514 18.9338V11.5081Z" fill="#0C0A93"/>
  <path d="M70.451 17.6224H73.4032C74.1745 17.6224 74.5602 17.9284 74.5602 18.5404V18.8846C74.5602 19.0376 74.4738 19.1141 74.3009 19.1141H68.7155C68.5692 19.1141 68.4961 19.054 68.4961 18.9338V10.1967C68.4961 10.0656 68.5692 10 68.7155 10H68.895C69.4004 10 69.786 10.1147 70.052 10.3442C70.318 10.5628 70.451 10.9289 70.451 11.4425V17.6224Z" fill="#0C0A93"/>
  <path d="M61.6643 19.2781C60.7999 19.2781 60.0951 19.1633 59.5499 18.9338C59.0046 18.7043 58.5857 18.4093 58.2932 18.0486C58.0006 17.688 57.8011 17.3055 57.6947 16.9012C57.5884 16.4968 57.5352 16.1253 57.5352 15.7865V10.1967C57.5352 10.0656 57.6083 10 57.7546 10H57.9142C58.4461 10 58.8384 10.1093 59.0911 10.3278C59.357 10.5464 59.49 10.9125 59.49 11.4261V15.475C59.49 16.1526 59.6496 16.7045 59.9688 17.1307C60.2879 17.5569 60.8531 17.77 61.6643 17.77C62.4888 17.77 63.0673 17.5569 63.3998 17.1307C63.7323 16.6935 63.8985 16.1417 63.8985 15.475V10.1967C63.8985 10.0656 63.9783 10 64.1379 10H64.2974C64.8161 10 65.2084 10.1093 65.4744 10.3278C65.7403 10.5464 65.8733 10.9125 65.8733 11.4261V15.7701C65.8733 16.087 65.8201 16.4476 65.7137 16.852C65.6073 17.2454 65.4079 17.6279 65.1153 17.9995C64.8227 18.371 64.3972 18.677 63.8386 18.9174C63.2934 19.1578 62.5686 19.2781 61.6643 19.2781Z" fill="#0C0A93"/>
  <path d="M48.7311 16.8192L47.8335 18.901C47.7803 19.0431 47.6805 19.1141 47.5342 19.1141H45.9983C45.9318 19.1141 45.8852 19.0977 45.8586 19.065C45.832 19.0212 45.8254 18.9775 45.8387 18.9338L49.7085 10.1639C49.7484 10.0546 49.8282 10 49.9479 10H50.6062C51.0849 10 51.4573 10.082 51.7233 10.2459C51.9892 10.4098 52.2286 10.7322 52.4414 11.213L55.8524 18.9338C55.879 18.9775 55.879 19.0212 55.8524 19.065C55.8258 19.0977 55.7793 19.1141 55.7128 19.1141H54.1369C53.9375 19.1141 53.8111 19.0431 53.7579 18.901L52.8603 16.8192H48.7311ZM49.3295 15.4586H52.3017L50.8057 11.9343L49.3295 15.4586Z" fill="#0C0A93"/>
  <path d="M40.8573 17.2126L43.5502 10.2295C43.6167 10.0765 43.743 10 43.9292 10H45.3655C45.5117 10 45.5649 10.0656 45.525 10.1967L41.9544 18.9338C41.9012 19.054 41.8148 19.1141 41.6951 19.1141H39.8798C39.7203 19.1141 39.6205 19.054 39.5806 18.9338L36.01 10.2131C35.9967 10.1585 35.9967 10.1093 36.01 10.0656C36.0366 10.0219 36.0831 10 36.1496 10H36.8079C37.26 10 37.6124 10.082 37.8651 10.2459C38.1311 10.3989 38.3572 10.7213 38.5433 11.213L40.8573 17.2126Z" fill="#0C0A93"/>
  <path fill-rule="evenodd" clip-rule="evenodd" d="M16.1978 29.1445L9.37721 18.2207L20.7534 0H34.3945L16.1978 29.1445ZM16.3682 7.02357L8.25561 0L0.00121689 7.02357H16.3682Z" fill="#0C0A93"/>
</svg>'''

VAULT_MARK_DATA_URI = "data:image/svg+xml," + "%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 35 30' fill='none'%3E%3Cpath fill-rule='evenodd' clip-rule='evenodd' d='M16.1978 29.1445L9.37721 18.2207L20.7534 0H34.3945L16.1978 29.1445ZM16.3682 7.02357L8.25561 0L0.00121689 7.02357H16.3682Z' fill='%230C0A93'/%3E%3C/svg%3E"

USER_AVATAR = "data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='1' height='1'/%3E"

SUGGESTION_CHIPS = [
    "What is E-Khata?",
    "Khata Transfer process",
    "My loan is closed. How do I cancel MODT in Bangalore?",
    "How long does it take to prepare a sale deed?",
    "I found a resale property. Can you check if it’s legally clear?",
    "Vault service pricing",
]


# ── Custom CSS ──────────────────────────────────────────────────────────────
st.markdown("""
<style>
/* ─── Font ─── */
@import url('https://fonts.googleapis.com/css2?family=DM+Sans:ital,opsz,wght@0,9..40,300;0,9..40,400;0,9..40,500;0,9..40,600;0,9..40,700&display=swap');

/* ═══════════════════════════════════════════════════════
   GLOBAL OVERRIDES
   ═══════════════════════════════════════════════════════ */
:root {
    --brand:        #0C0A93;
    --brand-light:  #4D4BFF;
    --brand-dim:    rgba(12,10,147,0.05);
    --brand-glow:   rgba(12,10,147,0.08);
    --text:         #111827;
    --text-2:       #4B5563;
    --text-3:       #9CA3AF;
    --border:       #E5E7EB;
    --surface:      #F9FAFB;
    --radius:       10px;
    --radius-lg:    14px;
    --ease:         cubic-bezier(.4,0,.2,1);
}

html, body, .stApp, .stApp > *, [class*="css"] {
    font-family: 'DM Sans', -apple-system, BlinkMacSystemFont, sans-serif !important;
    -webkit-font-smoothing: antialiased !important;
}

/* Hide default Streamlit headings (we use custom HTML) */
.main h1, .main [data-testid="stHeading"] {
    display: none !important;
    height: 0 !important;
    overflow: hidden !important;
    margin: 0 !important;
    padding: 0 !important;
}

section[data-testid="stSidebar"] h1 {
    display: none !important;
    height: 0 !important;
    margin: 0 !important;
    padding: 0 !important;
}

/* ═══════════════════════════════════════════════════════
   HIDE STREAMLIT CHROME (keep sidebar toggle)
   ═══════════════════════════════════════════════════════ */
.stDeployButton,
#MainMenu,
footer,
.stActionButton,
.styles_terminalButton__JBj5T,
[data-testid="stDecoration"],
#stDecoration,
.viewerBadge_container__r5tak,
.styles_viewerBadge__CvC9N {
    display: none !important;
    visibility: hidden !important;
}
/* Hide toolbar action buttons (Fork/GitHub) but keep sidebar toggle */
[data-testid="stToolbar"] > [data-testid="stToolbarActions"] {
    display: none !important;
}
/* Make header transparent but keep it interactive (sidebar toggle lives here) */
header[data-testid="stHeader"] {
    background: transparent !important;
    border-bottom: none !important;
}

/* ── Visit Vault sidebar button ── */
.visit-vault-sb a {
    display: flex;
    align-items: center;
    justify-content: center;
    gap: 0.5rem;
    padding: 0.6rem 1rem;
    background: linear-gradient(135deg, #0C0A93 0%, #4D4BFF 100%);
    color: #FFFFFF;
    border: none;
    border-radius: 10px;
    font-family: 'DM Sans', sans-serif;
    font-size: 0.82rem;
    font-weight: 600;
    text-decoration: none;
    transition: all 0.25s var(--ease);
    box-shadow: 0 2px 10px rgba(12,10,147,0.2);
    letter-spacing: 0.2px;
}
.visit-vault-sb a:hover {
    background: linear-gradient(135deg, #1a18b8 0%, #6866FF 100%);
    box-shadow: 0 4px 18px rgba(12,10,147,0.3);
    transform: translateY(-1px);
}
.visit-vault-sb a svg {
    width: 14px; height: 14px;
    stroke: #FFFFFF; stroke-width: 2;
    fill: none;
}

/* ═══════════════════════════════════════════════════════
   SIDEBAR
   ═══════════════════════════════════════════════════════ */
section[data-testid="stSidebar"] {
    background: #FAFBFC !important;
    border-right: 1px solid var(--border) !important;
}

section[data-testid="stSidebar"] > div {
    padding-top: 0 !important;
    padding-left: 0.8rem !important;
    padding-right: 0.8rem !important;
}

/* Remove any top spacing from first element */
section[data-testid="stSidebar"] [data-testid="stVerticalBlock"] {
    padding-top: 0 !important;
    margin-top: 0 !important;
}

/* Brand header — flush to edges and top */
.sb-header {
    background: linear-gradient(135deg, #0C0A93 0%, #1a18b8 100%);
    margin: 0 -0.8rem;
    margin-top: 0 !important;
    padding: 1.8rem 1.2rem 1.5rem 1.2rem;
}
.sb-header svg { height: 24px; width: auto; }
.sb-header svg path { fill: white !important; }
.sb-header svg path:last-child { fill: rgba(255,255,255,0.75) !important; }
.sb-header-label {
    font-size: 0.56rem; font-weight: 600;
    color: rgba(255,255,255,0.6);
    letter-spacing: 2.5px; text-transform: uppercase;
}

/* Section labels */
section[data-testid="stSidebar"] h2,
section[data-testid="stSidebar"] h3 {
    color: var(--text-3) !important;
    font-size: 0.6rem !important;
    text-transform: uppercase;
    letter-spacing: 1.5px;
    font-weight: 600 !important;
    margin-top: 1rem !important;
    margin-bottom: 0.5rem !important;
    padding-left: 0.2rem !important;
}

section[data-testid="stSidebar"] .stMarkdown p {
    color: var(--text-2) !important;
    font-size: 0.82rem !important;
}

section[data-testid="stSidebar"] hr {
    border-color: var(--border) !important;
    margin: 0.6rem 0 !important;
}

/* ── ALL sidebar buttons base ── */
section[data-testid="stSidebar"] button {
    background: #FFFFFF !important;
    color: var(--text-2) !important;
    border: 1px solid var(--border) !important;
    border-radius: 8px !important;
    font-size: 0.82rem !important;
    font-weight: 400 !important;
    padding: 0.5rem 0.7rem !important;
    transition: all 0.15s var(--ease) !important;
    text-align: left !important;
    box-shadow: 0 1px 2px rgba(0,0,0,0.03) !important;
    transform: none !important;
}

section[data-testid="stSidebar"] button:hover {
    background: #F3F4F6 !important;
    color: var(--text) !important;
    border-color: #D1D5DB !important;
}

section[data-testid="stSidebar"] button:active {
    background: #E5E7EB !important;
}

/* ── Active chat highlighting (#1) ── */
.chat-active button {
    background: var(--brand-dim) !important;
    border-color: var(--brand) !important;
    border-left: 3px solid var(--brand) !important;
    color: var(--brand) !important;
    font-weight: 600 !important;
}
.chat-active button:hover {
    background: rgba(12,10,147,0.08) !important;
    border-color: var(--brand) !important;
    color: var(--brand) !important;
}

/* ── New Chat button (scoped via wrapper div) ── */
.new-chat-btn button {
    background: var(--brand) !important;
    color: #FFFFFF !important;
    border: none !important;
    border-radius: 8px !important;
    font-weight: 600 !important;
    font-size: 0.82rem !important;
    padding: 0.5rem 1rem !important;
    box-shadow: 0 1px 6px rgba(12,10,147,0.18) !important;
    text-align: center !important;
    transition: all 0.25s var(--ease) !important;
}

.new-chat-btn button:hover {
    background: var(--brand-light) !important;
    color: #FFFFFF !important;
    box-shadow: 0 4px 16px rgba(12,10,147,0.25) !important;
    transform: translateY(-1px) !important;
}

.new-chat-btn button:active {
    background: var(--brand) !important;
    color: #FFFFFF !important;
    transform: scale(0.97) !important;
    box-shadow: 0 1px 3px rgba(12,10,147,0.12) !important;
}

/* ── Chat history row — force vertical alignment ── */
section[data-testid="stSidebar"] [data-testid="stHorizontalBlock"] {
    align-items: center !important;
    gap: 0.25rem !important;
    margin-top: -0.22rem !important;
    margin-bottom: -0.22rem !important;
}

/* Collapse Streamlit's default vertical block spacing in sidebar */
section[data-testid="stSidebar"] [data-testid="stVerticalBlock"] > [data-testid="stVerticalBlockBorderWrapper"] {
    margin-bottom: -0.22rem !important;
    padding-top: 0 !important;
    padding-bottom: 0 !important;
}
section[data-testid="stSidebar"] [data-testid="stVerticalBlock"] {
    gap: 0.3rem !important;
}

/* ── Delete button ── */
.del-btn { margin: 0; padding: 0; }
.del-btn button {
    width: 32px !important;
    height: 32px !important;
    min-width: 32px !important;
    padding: 0 !important;
    font-size: 0.72rem !important;
    text-align: center !important;
    color: #B0B4BA !important;
    border: 1px solid var(--border) !important;
    border-radius: 8px !important;
    background: #FFFFFF !important;
    box-shadow: 0 1px 2px rgba(0,0,0,0.03) !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
}
.del-btn button:hover {
    background: rgba(239,68,68,0.08) !important;
    color: #EF4444 !important;
    border-color: rgba(239,68,68,0.2) !important;
}
.del-btn button:active {
    background: rgba(239,68,68,0.15) !important;
    color: #DC2626 !important;
}

/* Delete confirm state (red) */
.del-confirm { margin: 0; padding: 0; }
.del-confirm button {
    width: auto !important;
    min-width: 32px !important;
    height: 32px !important;
    padding: 0 0.6rem !important;
    font-size: 0.68rem !important;
    font-weight: 600 !important;
    color: #FFFFFF !important;
    background: #EF4444 !important;
    border: none !important;
    border-radius: 8px !important;
    box-shadow: 0 1px 4px rgba(239,68,68,0.25) !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
    text-align: center !important;
    animation: confirm-pulse 0.3s var(--ease) !important;
}
.del-confirm button:hover {
    background: #DC2626 !important;
    color: #FFFFFF !important;
    box-shadow: 0 2px 8px rgba(239,68,68,0.35) !important;
}

@keyframes confirm-pulse {
    0%   { transform: scale(0.9); opacity: 0.7; }
    100% { transform: scale(1);   opacity: 1; }
}

/* ── Compact About / sidebar footer (#7) ── */
.sb-footer {
    font-size: 0.72rem;
    color: var(--text-3);
    padding: 0.6rem 0.3rem;
    line-height: 1.6;
}
.sb-footer strong { color: var(--text-2); font-weight: 600; }
.sb-footer .sf-row {
    display: flex; align-items: center; gap: 0.4rem;
    margin-top: 0.2rem; font-size: 0.7rem;
}
.sb-footer .sf-row svg {
    width: 12px; height: 12px; flex-shrink: 0; color: var(--brand); opacity: 0.6;
}

/* ═══════════════════════════════════════════════════════
   MAIN CONTENT
   ═══════════════════════════════════════════════════════ */
.main .block-container {
    max-width: 820px !important;
    padding-top: 2.5rem !important;
    padding-bottom: 0.5rem !important;
}

/* ── Hero header ── */
.hero-block {
    text-align: center;
    padding: 1.5rem 0 0.5rem 0;
}
.hero-block svg {
    height: 28px; width: auto; margin-bottom: 0.6rem;
}
.hero-block h2 {
    font-family: 'DM Sans', sans-serif;
    font-size: 1.5rem;
    font-weight: 700;
    color: var(--text);
    letter-spacing: -0.5px;
    margin: 0 0 0.3rem 0;
}
.hero-block p {
    color: var(--text-3);
    font-size: 0.9rem;
    margin: 0;
}

/* ── Empty state with background texture (#5) ── */
.empty-state {
    text-align: center;
    padding: 3rem 1rem 1rem 1rem;
    position: relative;
}
.empty-state::before {
    content: '';
    position: absolute;
    top: -1rem;
    left: 50%;
    transform: translateX(-50%);
    width: 500px;
    height: 500px;
    background: radial-gradient(circle, rgba(12,10,147,0.04) 0%, rgba(77,75,255,0.02) 40%, transparent 70%);
    border-radius: 50%;
    pointer-events: none;
    z-index: 0;
}
.empty-state > * {
    position: relative;
    z-index: 1;
}
.empty-state .empty-icon {
    width: 64px;
    height: 64px;
    margin: 0 auto 1.2rem auto;
    background: linear-gradient(135deg, rgba(12,10,147,0.08) 0%, rgba(77,75,255,0.06) 100%);
    border-radius: 20px;
    display: flex;
    align-items: center;
    justify-content: center;
}
.empty-state .empty-icon svg {
    height: 28px; width: auto; opacity: 0.6;
}
.empty-state h3 {
    font-size: 1.1rem;
    font-weight: 600;
    color: var(--text);
    margin: 0 0 0.4rem 0;
}
.empty-state p {
    font-size: 0.88rem;
    color: var(--text-3);
    max-width: 380px;
    margin: 0 auto 1.5rem auto;
    line-height: 1.6;
}

/* ── Suggestion chips (now styled as buttons) (#2) ── */
.chip-grid {
    display: flex;
    flex-wrap: wrap;
    gap: 0.5rem;
    justify-content: center;
    max-width: 520px;
    margin: 0 auto;
}
.chip-grid button {
    display: inline-flex !important;
    align-items: center !important;
    padding: 0.5rem 0.9rem !important;
    background: #FFFFFF !important;
    border: 1px solid var(--border) !important;
    border-radius: 100px !important;
    font-size: 0.78rem !important;
    color: var(--text-2) !important;
    cursor: pointer !important;
    transition: all 0.2s var(--ease) !important;
    box-shadow: 0 1px 2px rgba(0,0,0,0.03) !important;
    font-weight: 400 !important;
    text-align: center !important;
}
.chip-grid button:hover {
    border-color: var(--brand) !important;
    color: var(--brand) !important;
    box-shadow: 0 2px 8px rgba(12,10,147,0.08) !important;
    transform: translateY(-1px) !important;
}
.chip-grid button:active {
    transform: scale(0.96) !important;
    box-shadow: 0 0 0 2px rgba(12,10,147,0.08) !important;
}

/* ═══════════════════════════════════════════════════════
   CHAT MESSAGES
   ═══════════════════════════════════════════════════════ */
div[data-testid="stChatMessage"] {
    border-radius: var(--radius-lg) !important;
    padding: 1.2rem 1.5rem !important;
    margin-bottom: 1rem !important;
    transition: all 0.3s var(--ease) !important;
    animation: messageSlideIn 0.4s var(--ease) !important;
    opacity: 0;
    animation-fill-mode: forwards !important;
}

@keyframes messageSlideIn {
    from {
        opacity: 0;
        transform: translateY(10px);
    }
    to {
        opacity: 1;
        transform: translateY(0);
    }
}

/* User message - clean white with blue accent */
div[data-testid="stChatMessage"]:has([data-testid="chatAvatarIcon-user"]) {
    background: #FFFFFF !important;
    border: 1px solid var(--border) !important;
    border-left: 4px solid var(--brand) !important;
    box-shadow: 0 2px 8px rgba(12,10,147,0.08) !important;
}

div[data-testid="stChatMessage"]:has([data-testid="chatAvatarIcon-user"]):hover {
    box-shadow: 0 4px 12px rgba(12,10,147,0.12) !important;
    transform: translateX(2px) !important;
}

/* Assistant message - subtle background with gradient accent */
div[data-testid="stChatMessage"]:has([data-testid="chatAvatarIcon-assistant"]) {
    background: linear-gradient(135deg, rgba(12,10,147,0.02) 0%, rgba(77,75,255,0.03) 100%) !important;
    border: 1px solid rgba(12,10,147,0.08) !important;
    border-left: 4px solid var(--brand-light) !important;
    box-shadow: 0 2px 10px rgba(77,75,255,0.06) !important;
}

div[data-testid="stChatMessage"]:has([data-testid="chatAvatarIcon-assistant"]):hover {
    box-shadow: 0 4px 14px rgba(77,75,255,0.1) !important;
    transform: translateX(-2px) !important;
}

/* Avatar styling */
div[data-testid="stChatMessage"] img {
    width: 32px !important; 
    height: 32px !important; 
    border-radius: 8px !important;
    box-shadow: 0 2px 6px rgba(0,0,0,0.1) !important;
}

/* User avatar - add subtle border */
div[data-testid="stChatMessage"]:has([data-testid="chatAvatarIcon-user"]) img {
    border: 2px solid var(--brand) !important;
}

/* Assistant avatar - add subtle glow */
div[data-testid="stChatMessage"]:has([data-testid="chatAvatarIcon-assistant"]) img {
    border: 2px solid var(--brand-light) !important;
    box-shadow: 0 0 12px rgba(77,75,255,0.3) !important;
}

div[data-testid="stChatMessage"] p {
    color: var(--text) !important; line-height: 1.75 !important;
    font-size: 0.9rem !important; font-weight: 400 !important;
}

div[data-testid="stChatMessage"] strong { color: var(--text) !important; font-weight: 600 !important; }

div[data-testid="stChatMessage"] a {
    color: var(--brand) !important; text-decoration: none !important;
    font-weight: 500 !important;
    border-bottom: 1px solid rgba(12,10,147,0.2) !important;
    transition: all 0.2s var(--ease) !important;
}
div[data-testid="stChatMessage"] a:hover { border-bottom-color: var(--brand) !important; }

div[data-testid="stChatMessage"] li {
    color: var(--text) !important; font-size: 0.9rem !important; line-height: 1.75 !important;
}
div[data-testid="stChatMessage"] li::marker { color: var(--brand-light) !important; }

div[data-testid="stChatMessage"] h1,
div[data-testid="stChatMessage"] h2,
div[data-testid="stChatMessage"] h3 {
    display: block !important; color: var(--text) !important;
    font-weight: 700 !important; font-size: 1.05rem !important;
    margin-top: 1rem !important; margin-bottom: 0.35rem !important;
    letter-spacing: -0.2px !important;
}

/* Code blocks in messages */
div[data-testid="stChatMessage"] code {
    background: #F3F4F6 !important;
    color: var(--brand) !important;
    padding: 0.15rem 0.4rem !important;
    border-radius: 5px !important;
    font-size: 0.82rem !important;
}

/* ═══════════════════════════════════════════════════════
   CHAT INPUT — Floating style
   ═══════════════════════════════════════════════════════ */
div[data-testid="stChatInput"] {
    padding-bottom: 1rem !important;
}

div[data-testid="stBottom"] > div,
div[data-testid="stBottom"] {
    background: #FFFFFF !important;
}

div[data-testid="stChatInput"] > div {
    background: #FFFFFF !important;
    border: 1.5px solid var(--border) !important;
    border-radius: 16px !important;
    transition: all 0.3s var(--ease) !important;
    box-shadow: 0 2px 12px rgba(0,0,0,0.06), 0 0 0 0px transparent !important;
}

div[data-testid="stChatInput"] > div:focus-within {
    border-color: var(--brand) !important;
    box-shadow: 0 4px 20px rgba(12,10,147,0.08), 0 0 0 3px rgba(12,10,147,0.06) !important;
}

div[data-testid="stChatInput"] textarea {
    color: var(--text) !important;
    font-family: 'DM Sans', sans-serif !important;
    font-size: 0.92rem !important;
    box-shadow: none !important; border: none !important; background: transparent !important;
}
div[data-testid="stChatInput"] textarea::placeholder { color: #C9CCD1 !important; }

div[data-testid="stChatInput"] button {
    color: var(--brand) !important;
    transition: all 0.2s var(--ease) !important;
    border-radius: 10px !important;
}
div[data-testid="stChatInput"] button:hover {
    background: var(--brand-dim) !important;
    color: var(--brand-light) !important;
    transform: scale(1.1) !important;
}
div[data-testid="stChatInput"] button:active {
    transform: scale(0.92) !important;
    transition: all 0.08s !important;
}

/* ═══════════════════════════════════════════════════════
   TYPING INDICATOR (#4)
   ═══════════════════════════════════════════════════════ */
.typing-indicator {
    display: inline-flex;
    align-items: center;
    gap: 5px;
    padding: 1rem 0;
}
.typing-indicator span {
    width: 8px;
    height: 8px;
    background: linear-gradient(135deg, var(--brand) 0%, var(--brand-light) 100%);
    border-radius: 50%;
    animation: typing-bounce 1.4s infinite ease-in-out both;
    box-shadow: 0 2px 4px rgba(77,75,255,0.2);
}
.typing-indicator span:nth-child(1) { animation-delay: 0s; }
.typing-indicator span:nth-child(2) { animation-delay: 0.16s; }
.typing-indicator span:nth-child(3) { animation-delay: 0.32s; }

@keyframes typing-bounce {
    0%, 80%, 100% { 
        transform: scale(0.7) translateY(0); 
        opacity: 0.4; 
    }
    40% { 
        transform: scale(1.1) translateY(-8px); 
        opacity: 1; 
    }
}

/* ═══════════════════════════════════════════════════════
   SPINNER & ALERTS
   ═══════════════════════════════════════════════════════ */
.stSpinner > div { border-top-color: var(--brand) !important; }
.stSpinner p { color: var(--text-2) !important; font-size: 0.82rem !important; }

div[data-testid="stAlert"] {
    background: rgba(239,68,68,0.04) !important;
    border: 1px solid rgba(239,68,68,0.12) !important;
    border-radius: var(--radius) !important;
    color: var(--text) !important;
}

/* ═══════════════════════════════════════════════════════
   SCROLLBAR
   ═══════════════════════════════════════════════════════ */
::-webkit-scrollbar { width: 6px; }
::-webkit-scrollbar-track { background: transparent; }
::-webkit-scrollbar-thumb { background: rgba(0,0,0,0.08); border-radius: 3px; }
::-webkit-scrollbar-thumb:hover { background: rgba(0,0,0,0.15); }

/* ═══════════════════════════════════════════════════════
   FOOTER
   ═══════════════════════════════════════════════════════ */
.vault-footer {
    text-align: center;
    padding: 2rem 1rem 1rem 1rem;
    max-width: 720px;
    margin: 0 auto;
}
.vault-footer .vf-copyright {
    font-size: 0.75rem;
    color: var(--text-3);
    letter-spacing: 0.3px;
    margin-bottom: 1rem;
}
.vault-footer .vf-brand {
    color: var(--brand);
    font-weight: 600;
}
.vault-footer .vf-sep {
    margin: 0 0.5rem;
    opacity: 0.3;
}
.vault-footer .vf-disclaimer {
    font-size: 0.7rem;
    color: var(--text-3);
    line-height: 1.6;
    padding: 0.8rem 1rem;
    background: rgba(12,10,147,0.02);
    border: 1px solid rgba(12,10,147,0.08);
    border-radius: 8px;
    text-align: center;
}
.vault-footer .vf-disclaimer strong {
    color: var(--text-2);
    font-weight: 600;
}

.main hr { border-color: var(--border) !important; margin-top: 1rem !important; margin-bottom: 0 !important; }

/* ═══════════════════════════════════════════════════════
   TOAST
   ═══════════════════════════════════════════════════════ */
div[data-testid="stToast"] {
    background: var(--text) !important; color: #FFFFFF !important;
    border-radius: 12px !important;
    box-shadow: 0 12px 40px rgba(0,0,0,0.12) !important;
    font-size: 0.84rem !important; font-weight: 500 !important;
}

/* ═══════════════════════════════════════════════════════
   LOGIN PAGE
   ═══════════════════════════════════════════════════════ */
.login-page {
    display: flex;
    flex-direction: column;
    align-items: center;
    justify-content: center;
    min-height: 70vh;
    text-align: center;
}
.login-card {
    background: #FFFFFF;
    border: 1px solid var(--border);
    border-radius: 20px;
    padding: 2.5rem 3rem;
    max-width: 420px;
    width: 100%;
    box-shadow: 0 8px 30px rgba(0,0,0,0.06);
}
.login-card svg { height: 36px; width: auto; margin-bottom: 1rem; }
.login-card h2 {
    font-size: 1.3rem; font-weight: 700; color: var(--text);
    margin: 0.5rem 0 0.3rem 0; letter-spacing: -0.3px;
}
.login-card p {
    font-size: 0.88rem; color: var(--text-3); margin: 0 0 1.5rem 0; line-height: 1.6;
}
.login-card .google-btn {
    display: inline-flex; align-items: center; gap: 0.6rem;
    padding: 0.7rem 1.5rem;
    background: #FFFFFF;
    border: 1.5px solid var(--border);
    border-radius: 100px;
    font-size: 0.88rem; font-weight: 600;
    color: var(--text);
    text-decoration: none;
    transition: all 0.2s var(--ease);
    box-shadow: 0 2px 8px rgba(0,0,0,0.04);
}
.login-card .google-btn:hover {
    border-color: #4285F4;
    box-shadow: 0 4px 16px rgba(66,133,244,0.15);
    transform: translateY(-1px);
}
.login-card .google-btn img { width: 18px; height: 18px; }
.login-footer {
    font-size: 0.7rem; color: var(--text-3);
    margin-top: 1.2rem; line-height: 1.5;
}

/* ── User greeting in sidebar ── */
.user-greeting {
    display: flex; align-items: center; gap: 0.5rem;
    padding: 0.6rem 0.3rem;
    font-size: 0.8rem; color: var(--text-2);
    margin-top: 1.0rem;
    margin-bottom: 1.25rem;
}
.user-greeting img {
    width: 28px; height: 28px; border-radius: 50%;
    border: 1.5px solid var(--border);
}
.user-greeting .ug-name { font-weight: 600; color: var(--text); }
.user-greeting .ug-email { font-size: 0.68rem; color: var(--text-3); }

.logout-btn button {
    background: transparent !important;
    color: var(--text-3) !important;
    border: 1px solid var(--border) !important;
    font-size: 0.72rem !important;
    padding: 0.3rem 0.6rem !important;
    border-radius: 6px !important;
}
.logout-btn button:hover {
    background: rgba(239,68,68,0.06) !important;
    color: #EF4444 !important;
    border-color: rgba(239,68,68,0.2) !important;
}

/* ═══════════════════════════════════════════════════════
   DRAFT CANVAS — Isolated document panel
   ═══════════════════════════════════════════════════════ */
.draft-canvas {
    border: 2px solid var(--brand);
    border-radius: var(--radius-lg);
    margin: 1.2rem 0;
    overflow: hidden;
    box-shadow: 0 4px 20px rgba(12,10,147,0.08);
    animation: canvasSlideIn 0.5s var(--ease);
}
@keyframes canvasSlideIn {
    from { opacity: 0; transform: translateY(20px); }
    to { opacity: 1; transform: translateY(0); }
}
.draft-canvas-header {
    background: linear-gradient(135deg, #0C0A93 0%, #1a18b8 100%);
    padding: 0.8rem 1.2rem;
    display: flex;
    align-items: center;
    justify-content: space-between;
}
.draft-canvas-header .deed-badge {
    background: rgba(255,255,255,0.15);
    color: #FFFFFF;
    padding: 0.3rem 0.8rem;
    border-radius: 100px;
    font-size: 0.72rem;
    font-weight: 600;
    text-transform: uppercase;
    letter-spacing: 1px;
}
.draft-canvas-header .draft-meta {
    color: rgba(255,255,255,0.6);
    font-size: 0.68rem;
}
.draft-canvas-content {
    padding: 2rem 2.5rem;
    background: #FFFFFF;
    max-height: 70vh;
    overflow-y: auto;
    line-height: 1.8;
    font-size: 0.92rem;
    color: var(--text);
}
.draft-canvas-content h1, .draft-canvas-content h2, .draft-canvas-content h3 {
    color: var(--text) !important;
    display: block !important;
    font-weight: 700 !important;
    margin-top: 1.2rem !important;
    margin-bottom: 0.5rem !important;
}
.draft-canvas-content h1 { font-size: 1.2rem !important; }
.draft-canvas-content h2 { font-size: 1.05rem !important; }
.draft-canvas-content h3 { font-size: 0.95rem !important; }
.draft-canvas-actions {
    background: var(--surface);
    border-top: 1px solid var(--border);
    padding: 0.6rem 1.2rem;
    display: flex;
    align-items: center;
    gap: 0.6rem;
}
/* Export / download button styling */
.export-btn button, .download-btn button {
    background: linear-gradient(135deg, #0C0A93 0%, #4D4BFF 100%) !important;
    color: #FFFFFF !important;
    border: none !important;
    border-radius: 8px !important;
    font-weight: 600 !important;
    font-size: 0.78rem !important;
    padding: 0.4rem 1rem !important;
    box-shadow: 0 2px 8px rgba(12,10,147,0.2) !important;
    transition: all 0.25s var(--ease) !important;
}
.export-btn button:hover, .download-btn button:hover {
    background: linear-gradient(135deg, #1a18b8 0%, #6866FF 100%) !important;
    box-shadow: 0 4px 14px rgba(12,10,147,0.3) !important;
    transform: translateY(-1px) !important;
}

/* ── File Upload Section ── */
.file-upload-section {
    margin: 0.8rem 0;
    padding: 0.8rem;
    background: var(--surface);
    border: 1px dashed var(--border);
    border-radius: var(--radius);
}
.file-upload-section .upload-label {
    font-size: 0.72rem;
    color: var(--text-3);
    font-weight: 600;
    text-transform: uppercase;
    letter-spacing: 1px;
    margin-bottom: 0.4rem;
}
.uploaded-files-list {
    font-size: 0.78rem;
    color: var(--text-2);
}
.uploaded-files-list .file-item {
    display: flex;
    align-items: center;
    gap: 0.4rem;
    padding: 0.25rem 0;
}
.uploaded-files-list .file-item a {
    color: var(--brand) !important;
    text-decoration: none;
    font-weight: 500;
}
/* ═══════════════════════════════════════════════════════
   FREE USER CTA BANNER & LOCKED CANVAS
   ═══════════════════════════════════════════════════════ */
.free-cta-bar {
    display: flex;
    align-items: center;
    justify-content: center;
    gap: 0.8rem;
    padding: 0.7rem 1.2rem;
    margin: 0 auto 1rem auto;
    max-width: 720px;
    background: linear-gradient(135deg, rgba(12,10,147,0.04) 0%, rgba(77,75,255,0.06) 100%);
    border: 1px solid rgba(12,10,147,0.12);
    border-radius: 12px;
    font-size: 0.82rem;
    color: var(--text-2);
    animation: ctaFadeIn 0.6s var(--ease);
}
@keyframes ctaFadeIn {
    from { opacity: 0; transform: translateY(-8px); }
    to   { opacity: 1; transform: translateY(0); }
}
.cta-login-btn {
    display: inline-flex;
    align-items: center;
    gap: 0.4rem;
    padding: 0.35rem 1rem;
    background: linear-gradient(135deg, #0C0A93 0%, #4D4BFF 100%);
    color: #FFFFFF !important;
    border-radius: 100px;
    font-size: 0.78rem;
    font-weight: 600;
    text-decoration: none;
    transition: all 0.25s var(--ease);
    box-shadow: 0 2px 8px rgba(12,10,147,0.2);
    white-space: nowrap;
}
.cta-login-btn:hover {
    background: linear-gradient(135deg, #1a18b8 0%, #6866FF 100%);
    box-shadow: 0 4px 14px rgba(12,10,147,0.3);
    transform: translateY(-1px);
}
.cta-login-btn img {
    width: 14px;
    height: 14px;
}
.draft-canvas-actions.locked {
    justify-content: center;
    padding: 0.8rem 1.2rem;
    font-size: 0.82rem;
    color: var(--text-2);
    gap: 0.8rem;
}
</style>
""", unsafe_allow_html=True)

# Initialize session state
if "messages" not in st.session_state:
    st.session_state.messages = []
if "chat_id" not in st.session_state:
    st.session_state.chat_id = str(uuid.uuid4())
if "chat_name" not in st.session_state:
    st.session_state.chat_name = "New Chat"
if "drafting_mode" not in st.session_state:
    st.session_state.drafting_mode = False
if "current_deed_type" not in st.session_state:
    st.session_state.current_deed_type = ""
if "draft_content" not in st.session_state:
    st.session_state.draft_content = ""
if "uploaded_docs_context" not in st.session_state:
    st.session_state.uploaded_docs_context = ""
if "uploaded_files_info" not in st.session_state:
    st.session_state.uploaded_files_info = []

# ═══════════════════════════════════════════════════════
# AUTH CHECK — Soft gate (anonymous users can still chat)
# ═══════════════════════════════════════════════════════
user = check_auth()
is_logged_in = user is not None

# Generate login URL (needed for both CTA banner and sidebar)
client_id, client_secret, redirect_uri = get_auth_config()
login_url = get_login_url(client_id, redirect_uri) if client_id else "#"

# ── Derive user info safely ──
if is_logged_in:
    user_email = user.get("email", "")
    user_name = user.get("name", "User")
    user_picture = user.get("picture", "")

    # On-login migration: push any anonymous session messages to Firebase
    if st.session_state.get("_was_anonymous") and st.session_state.messages:
        save_chat(user_email, user_name, st.session_state.chat_id,
                  st.session_state.messages, st.session_state.chat_name)
        if st.session_state.draft_content:
            save_draft(user_email, user_name, st.session_state.chat_id,
                       st.session_state.draft_content,
                       st.session_state.current_deed_type,
                       "Migrated from free session",
                       st.session_state.chat_name)
        st.session_state.pop("_was_anonymous", None)
        st.toast("✅ Your conversation has been saved!")
else:
    user_email = ""
    user_name = ""
    user_picture = ""
    # Mark this session as anonymous so we can migrate on login
    if "_was_anonymous" not in st.session_state:
        st.session_state["_was_anonymous"] = True

# Sidebar — logged-in users only
if is_logged_in:
    with st.sidebar:
        st.title("Vault PropTech")  # Hidden by CSS

        # Brand header — indigo banner at top
        st.markdown(f"""
        <div class="sb-header">
            {VAULT_LOGO_SVG}
            <div class="sb-header-label">Legal Assistant</div>
        </div>
        """, unsafe_allow_html=True)

        # User greeting
        st.markdown(f"""
        <div class="user-greeting">
            <img src="{user_picture}" alt="" referrerpolicy="no-referrer">
            <div>
                <div class="ug-name">{user_name}</div>
                <div class="ug-email">{user_email}</div>
            </div>
        </div>
        """, unsafe_allow_html=True)

        st.markdown('<div class="new-chat-btn">', unsafe_allow_html=True)
        if st.button("New Chat", use_container_width=True):
            st.session_state.messages = []
            st.session_state.chat_id = str(uuid.uuid4())
            st.session_state.chat_name = "New Chat"
            st.session_state.drafting_mode = False
            st.session_state.current_deed_type = ""
            st.session_state.draft_content = ""
            st.session_state.uploaded_docs_context = ""
            st.session_state.uploaded_files_info = []
            st.toast("New conversation started")
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown("---")
        st.subheader("Chat History")

        chats = load_chats(user_email)

        if not chats:
            st.markdown(
                "<p style='color:#9CA3AF; font-size:0.78rem; padding:0.3rem 0; font-style:italic;'>No conversations yet</p>",
                unsafe_allow_html=True,
            )

        for chat in chats:
            is_active = chat["chat_id"] == st.session_state.chat_id
            col1, col2 = st.columns([4, 1])
            with col1:
                if is_active:
                    st.markdown('<div class="chat-active">', unsafe_allow_html=True)
                if st.button(chat["chat_name"], key=chat["chat_id"], use_container_width=True):
                    loaded = load_chat(user_email, chat["chat_id"])
                    if loaded:
                        st.session_state.messages = loaded["messages"]
                        st.session_state.chat_id = chat["chat_id"]
                        st.session_state.chat_name = chat["chat_name"]
                        # Restore draft state if this chat has a draft
                        draft_msg = None
                        for m in loaded["messages"]:
                            if m.get("role") == "draft":
                                draft_msg = m
                                break
                        if draft_msg:
                            st.session_state.drafting_mode = True
                            st.session_state.draft_content = draft_msg.get("content", "")
                            st.session_state.current_deed_type = draft_msg.get("deed_type", "")
                        else:
                            st.session_state.drafting_mode = False
                            st.session_state.draft_content = ""
                            st.session_state.current_deed_type = ""
                        st.rerun()
                if is_active:
                    st.markdown('</div>', unsafe_allow_html=True)
            with col2:
                st.markdown('<div class="del-btn">', unsafe_allow_html=True)
                if st.button("\u2715", key=f"del_{chat['chat_id']}"):
                    delete_chat(user_email, chat["chat_id"])
                    if chat["chat_id"] == st.session_state.chat_id:
                        st.session_state.messages = []
                        st.session_state.chat_id = str(uuid.uuid4())
                        st.session_state.chat_name = "New Chat"
                    st.toast("Chat deleted")
                    st.rerun()
                st.markdown('</div>', unsafe_allow_html=True)

        # Compact sidebar footer
        st.markdown("---")
        st.markdown("""
        <div class="sb-footer">
            <strong>Vault PropTech</strong> · Property legal services, Bengaluru
            <div class="sf-row">
                <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M22 16.92v3a2 2 0 01-2.18 2 19.79 19.79 0 01-8.63-3.07 19.5 19.5 0 01-6-6 19.79 19.79 0 01-3.07-8.67A2 2 0 014.11 2h3a2 2 0 012 1.72c.127.96.361 1.903.7 2.81a2 2 0 01-.45 2.11L8.09 9.91a16 16 0 006 6l1.27-1.27a2 2 0 012.11-.45c.907.339 1.85.573 2.81.7A2 2 0 0122 16.92z"/></svg>
                +91 88619 50376
            </div>
            <div class="sf-row">
                <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M4 4h16c1.1 0 2 .9 2 2v12c0 1.1-.9 2-2 2H4c-1.1 0-2-.9-2-2V6c0-1.1.9-2 2-2z"/><polyline points="22,6 12,13 2,6"/></svg>
                info@vaultproptech.com
            </div>
        </div>
        """, unsafe_allow_html=True)

        st.markdown('<div style="margin-top: 1.2rem;"></div>', unsafe_allow_html=True)

        # Visit Vault button — appealing gradient style
        st.markdown("""
        <div class="visit-vault-sb">
            <a href="https://www.vaultproptech.com/" target="_blank" rel="noopener noreferrer">
                <svg viewBox="0 0 24 24"><path d="M18 13v6a2 2 0 01-2 2H5a2 2 0 01-2-2V8a2 2 0 012-2h6"/><polyline points="15 3 21 3 21 9"/><line x1="10" y1="14" x2="21" y2="3"/></svg>
                Visit Vault
            </a>
        </div>
        """, unsafe_allow_html=True)

        # Spacer
        st.markdown('<div style="margin-top: 1.8rem;"></div>', unsafe_allow_html=True)

        # Show uploaded files count if any (compact indicator in sidebar)
        if st.session_state.uploaded_files_info:
            count = len(st.session_state.uploaded_files_info)
            st.markdown("---")
            st.markdown(
                f'<p style="color:var(--text-3); font-size:0.72rem; font-weight:600;">'
                f'📎 {count} document{"s" if count > 1 else ""} attached</p>',
                unsafe_allow_html=True,
            )
            for f_info in st.session_state.uploaded_files_info:
                link = f_info.get("web_view_link", "")
                name = f_info.get("filename", "file")
                if link:
                    st.markdown(
                        f'<div class="file-item">📄 <a href="{link}" target="_blank" '
                        f'style="color:var(--brand);text-decoration:none;font-size:0.72rem;">{name}</a></div>',
                        unsafe_allow_html=True,
                    )
                else:
                    st.markdown(
                        f'<div class="file-item" style="font-size:0.72rem;">📄 {name}</div>',
                        unsafe_allow_html=True,
                    )

        # Logout button
        st.markdown('<div class="logout-btn">', unsafe_allow_html=True)
        if st.button("Sign Out", use_container_width=True):
            st.session_state.clear_cookie_flag = True
            st.session_state.user = None
            st.session_state.messages = []
            st.session_state.chat_id = str(uuid.uuid4())
            st.session_state.chat_name = "New Chat"
            st.session_state.drafting_mode = False
            st.session_state.draft_content = ""
            st.session_state.current_deed_type = ""
            st.session_state.uploaded_docs_context = ""
            st.session_state.uploaded_files_info = []
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

# Main content

# Hero — centered brand presentation
st.markdown(f"""
<div class="hero-block">
    {VAULT_LOGO_SVG}
    <h2>Legal Assistant</h2>
    <p>Property documentation & legal guidance · Exclusively in Bengaluru</p>
</div>
""", unsafe_allow_html=True)

# Free-user CTA banner
if not is_logged_in:
    st.markdown(f"""
    <div class="free-cta-bar">
        <b>SIGN IN</b>to save chats, export drafts, and upload documents.
        <a href="{login_url}" class="cta-login-btn">
            <img src="https://developers.google.com/identity/images/g-logo.png" alt="">
            Sign In
        </a>
    </div>
    """, unsafe_allow_html=True)

# Empty state with CLICKABLE suggestion chips (#2)
if not st.session_state.messages:
    st.markdown(f"""
    <div class="empty-state">
        <h3>How can I help you today?</h3>
        <p>Ask about property documentation, legal processes, or Vault's services in Bengaluru.</p>
    </div>
    """, unsafe_allow_html=True)

    # Clickable chips — these are real st.button elements
    st.markdown('<div class="chip-grid">', unsafe_allow_html=True)
    chip_cols = st.columns(3)
    for i, chip_text in enumerate(SUGGESTION_CHIPS):
        with chip_cols[i % 3]:
            if st.button(chip_text, key=f"chip_{i}"):
                st.session_state.chat_name = chip_text[:20] + ("..." if len(chip_text) > 20 else "")
                st.session_state.messages.append({"role": "user", "content": chip_text})
                st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════
# HELPERS — Generate downloadable PDF / DOCX from draft
# ═══════════════════════════════════════════════════════

def _strip_md(text: str) -> str:
    """Return plain text from Markdown (best-effort)."""
    text = re.sub(r'#{1,6}\s*', '', text)          # headings
    text = re.sub(r'\*{1,2}(.+?)\*{1,2}', r'\1', text)  # bold/italic
    text = re.sub(r'[-*]\s+', '- ', text)            # bullets
    text = text.encode('latin-1', errors='replace').decode('latin-1')  # safe for PDF
    return text.strip()


def _generate_pdf(content: str, deed_type: str) -> bytes:
    """Create a simple PDF from Markdown draft content."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_font('Helvetica', 'B', 16)
    title = deed_type.replace('_', ' ').title()
    pdf.cell(0, 12, title, ln=True, align='C')
    pdf.ln(6)
    pdf.set_font('Helvetica', '', 11)
    plain = _strip_md(content)
    for line in plain.split('\n'):
        stripped = line.strip()
        if not stripped:
            pdf.ln(4)
            continue
        pdf.multi_cell(0, 6, stripped)
        pdf.ln(2)
    return bytes(pdf.output())


def _generate_docx(content: str, deed_type: str) -> bytes:
    """Create a DOCX from Markdown draft content."""
    doc = DocxDocument()
    title = deed_type.replace('_', ' ').title()
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    plain = _strip_md(content)
    for line in plain.split('\n'):
        stripped = line.strip()
        if not stripped:
            continue
        p = doc.add_paragraph(stripped)
        for run in p.runs:
            run.font.size = Pt(11)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════
# DRAFT CANVAS — Render if a draft exists
# ═══════════════════════════════════════════════════════
if st.session_state.draft_content:
    deed_display = st.session_state.current_deed_type.replace('_', ' ').title()
    st.markdown(f"""
    <div class="draft-canvas">
        <div class="draft-canvas-header">
            <span class="deed-badge">📜 {deed_display}</span>
            <span class="draft-meta">Live Draft • Editable</span>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # Render draft content as markdown inside a styled container
    with st.container():
        st.markdown(
            f'<div class="draft-canvas-content">', unsafe_allow_html=True
        )
        st.markdown(st.session_state.draft_content)
        st.markdown('</div>', unsafe_allow_html=True)

    # Action buttons — logged-in users get full controls, free users see locked prompt
    if is_logged_in:
        st.markdown('<div class="draft-canvas-actions">', unsafe_allow_html=True)
        col_pdf, col_word, col_export, col_link = st.columns([1, 1, 1, 2])

        # ── Download as PDF ──
        with col_pdf:
            st.markdown('<div class="download-btn">', unsafe_allow_html=True)
            pdf_bytes = _generate_pdf(
                st.session_state.draft_content,
                st.session_state.current_deed_type,
            )
            deed_slug = st.session_state.current_deed_type.replace(' ', '_')
            st.download_button(
                label="⬇ Download PDF",
                data=pdf_bytes,
                file_name=f"Vault_Draft_{deed_slug}.pdf",
                mime="application/pdf",
                key="dl_pdf",
            )
            st.markdown('</div>', unsafe_allow_html=True)

        # ── Download as Word ──
        with col_word:
            st.markdown('<div class="download-btn">', unsafe_allow_html=True)
            docx_bytes = _generate_docx(
                st.session_state.draft_content,
                st.session_state.current_deed_type,
            )
            st.download_button(
                label="⬇ Download Word",
                data=docx_bytes,
                file_name=f"Vault_Draft_{deed_slug}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_docx",
            )
            st.markdown('</div>', unsafe_allow_html=True)

        # ── Export to Google Doc ──
        with col_export:
            st.markdown('<div class="export-btn">', unsafe_allow_html=True)
            if st.button("📄 Export to Google Doc", key="export_doc"):
                user_token = st.session_state.get("user_token", "")
                if not user_token:
                    st.warning("Please sign out and sign in again to enable Google Docs export.")
                else:
                    deed_display_name = st.session_state.current_deed_type.replace('_', ' ').title()
                    doc_title = f"Vault Draft - {deed_display_name}"
                    try:
                        with st.spinner("Creating Google Doc..."):
                            doc_url = export_to_google_doc(
                                doc_title,
                                st.session_state.draft_content,
                                user_token,
                            )
                        st.session_state.exported_doc_url = doc_url
                        save_draft(
                            user_email, user_name, st.session_state.chat_id,
                            st.session_state.draft_content,
                            st.session_state.current_deed_type,
                            f"Exported to Google Doc",
                            st.session_state.chat_name,
                            doc_link=doc_url,
                        )
                        st.toast("✅ Google Doc created!")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Export failed: {e}")
            st.markdown('</div>', unsafe_allow_html=True)

        # Show the Google Doc link if one exists
        with col_link:
            doc_url = st.session_state.get("exported_doc_url", "")
            if doc_url:
                st.link_button("🔗 Open Google Doc", doc_url, type="primary")
        st.markdown('</div>', unsafe_allow_html=True)
    else:
        # Free user — locked actions with sign-in prompt
        st.markdown(f"""
        <div class="draft-canvas-actions locked">
            🔒 Sign in to download or export this draft.
            <a href="{login_url}" class="cta-login-btn">
                <img src="https://developers.google.com/identity/images/g-logo.png" alt="">
                Sign In
            </a>
        </div>
        """, unsafe_allow_html=True)

# Display chat messages (skip draft messages — they show in the canvas)
for message in st.session_state.messages:
    if message.get("role") == "draft":
        continue  # Rendered in canvas above
    avatar = VAULT_MARK_DATA_URI if message["role"] == "assistant" else (user_picture or None)
    with st.chat_message(message["role"], avatar=avatar):
        st.markdown(message["content"])

# ═══════════════════════════════════════════════════════
# RESPONSE GENERATION — Handles both standard and drafting flows
# ═══════════════════════════════════════════════════════

def _handle_response(prompt_text):
    """Process a user message and generate a response (standard or drafting)."""
    kb_dir = os.path.dirname(os.path.abspath(__file__))

    result = ask(
        kb_dir, prompt_text,
        st.session_state.messages,
        user_name=user_name,
        user_email=user_email,
        uploaded_docs_context=st.session_state.uploaded_docs_context,
        existing_draft=st.session_state.draft_content,
        is_drafting_active=st.session_state.drafting_mode,
        verbose=False,
    )

    if isinstance(result, dict):
        # Drafting response — update canvas and save draft
        st.session_state.drafting_mode = True
        st.session_state.draft_content = result["draft"]
        st.session_state.current_deed_type = result["deed_type"]

        # Save draft to Firebase (logged-in only)
        if is_logged_in:
            save_draft(
                user_email, user_name, st.session_state.chat_id,
                result["draft"],
                result["deed_type"],
                result["summary"],
                st.session_state.chat_name,
            )

        # Add the short assistant message to chat
        assistant_msg = result["assistant_message"]
        st.session_state.messages.append({"role": "assistant", "content": assistant_msg})
        if is_logged_in:
            save_chat(user_email, user_name, st.session_state.chat_id,
                      st.session_state.messages, st.session_state.chat_name)

        return assistant_msg
    else:
        # Standard response
        st.session_state.messages.append({"role": "assistant", "content": result})
        if is_logged_in:
            save_chat(user_email, user_name, st.session_state.chat_id,
                      st.session_state.messages, st.session_state.chat_name)
        return result


# If last message is user (just asked), auto-trigger the response
if st.session_state.messages and st.session_state.messages[-1]["role"] == "user":
    # Check there's no assistant response right after
    needs_response = True
    if len(st.session_state.messages) >= 2 and st.session_state.messages[-1]["role"] != "user":
        needs_response = False

    if needs_response:
        last_prompt = st.session_state.messages[-1]["content"]
        with st.chat_message("assistant", avatar=VAULT_MARK_DATA_URI):
            typing_placeholder = st.empty()
            typing_placeholder.markdown(
                '<div class="typing-indicator"><span></span><span></span><span></span></div>',
                unsafe_allow_html=True
            )
            try:
                answer = _handle_response(last_prompt)
                typing_placeholder.empty()
                st.markdown(answer)
                # If drafting, rerun to show the canvas
                if st.session_state.drafting_mode:
                    st.rerun()
            except Exception as e:
                typing_placeholder.empty()
                error_msg = f"I encountered an error: {str(e)}\n\nPlease make sure:\n1. The knowledge base is ingested (`python setup.py ingest`)\n2. Your GEMINI_API_KEY is set in .env"
                st.error(error_msg)
                st.session_state.messages.append({"role": "assistant", "content": error_msg})

# ═══════════════════════════════════════════════════════
# FILE UPLOAD — logged-in users only
# ═══════════════════════════════════════════════════════
if is_logged_in:
    with st.expander("📎 Attach documents (PDF / Image)", expanded=False):
        st.markdown(
            '<p style="color:#9CA3AF; font-size:0.78rem; margin:0 0 0.5rem 0;">'
            'Upload previous deeds, Aadhaar/PAN, E-Khata, or property tax receipts. '
            'Uploading a document will enter legal drafting mode.</p>',
            unsafe_allow_html=True,
        )
        uploaded_file = st.file_uploader(
            "Upload PDF or Image",
            type=["pdf", "png", "jpg", "jpeg", "tiff", "bmp"],
            key="doc_uploader",
            label_visibility="collapsed",
        )
        if uploaded_file is not None:
            file_bytes = uploaded_file.read()
            mime_type = uploaded_file.type or "application/octet-stream"

            # Auto-enter drafting mode on file upload
            if not st.session_state.drafting_mode:
                st.session_state.drafting_mode = True
                if not st.session_state.current_deed_type:
                    st.session_state.current_deed_type = "legal_document"

            deed_name = st.session_state.current_deed_type or "legal_document"

            with st.spinner("Uploading to Drive..."):
                try:
                    result = upload_file(
                        user_email, deed_name,
                        file_bytes, uploaded_file.name, mime_type
                    )
                    st.toast(f"✅ Uploaded: {uploaded_file.name}")
                    st.session_state.uploaded_files_info.append(result)
                except Exception as e:
                    st.error(f"Upload failed: {e}")
                    result = None

            # Extract text using Gemini 2.5 Flash OCR
            if result:
                with st.spinner("Reading document with Gemini OCR..."):
                    extracted = extract_text_with_gemini(file_bytes, mime_type, uploaded_file.name)
                    st.session_state.uploaded_docs_context += (
                        f"\n\n--- {uploaded_file.name} ---\n{extracted}\n"
                    )
                st.rerun()

        # Show uploaded files inline
        if st.session_state.uploaded_files_info:
            for f_info in st.session_state.uploaded_files_info:
                link = f_info.get("web_view_link", "")
                name = f_info.get("filename", "file")
                if link:
                    st.markdown(f"📄 [{name}]({link})")
                else:
                    st.markdown(f"📄 {name}")

# Chat input
if prompt := st.chat_input("Ask your question here..."):
    if st.session_state.chat_name == "New Chat":
        st.session_state.chat_name = prompt[:20] + ("..." if len(prompt) > 20 else "")

    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user", avatar=(user_picture or None)):
        st.markdown(prompt)

    with st.chat_message("assistant", avatar=VAULT_MARK_DATA_URI):
        typing_placeholder = st.empty()
        typing_placeholder.markdown(
            '<div class="typing-indicator"><span></span><span></span><span></span></div>',
            unsafe_allow_html=True
        )

        try:
            answer = _handle_response(prompt)
            typing_placeholder.empty()
            st.markdown(answer)
            # If drafting, rerun to show/update the canvas
            if st.session_state.drafting_mode:
                st.rerun()
        except Exception as e:
            typing_placeholder.empty()
            error_msg = f"I encountered an error: {str(e)}\n\nPlease make sure:\n1. The knowledge base is ingested (`python setup.py ingest`)\n2. Your GEMINI_API_KEY is set in .env"
            st.error(error_msg)
            st.session_state.messages.append({"role": "assistant", "content": error_msg})


# Footer
st.markdown("---")
st.markdown("""
<div class='vault-footer'>
    <div class='vf-copyright'>
        © 2026 <span class='vf-brand'>Vault PropTech</span>
        <span class='vf-sep'>·</span>
        Bengaluru, Karnataka
    </div>
    <div class='vf-disclaimer'>
        <strong>Disclaimer:</strong> This platform provides information for informational purposes only and 
        does not constitute professional legal advice.
        <br>
        For accurate legal guidance tailored to your situation, 
        please consult a qualified legal professional - preferably through Vault.
    </div>
</div>
""", unsafe_allow_html=True)